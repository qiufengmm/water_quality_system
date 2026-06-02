"""Generate Week 1 progress report Word document."""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

SECTION = doc.sections[0]
SECTION.page_width = Emu(7560310)
SECTION.page_height = Emu(10692130)
SECTION.top_margin = Cm(2.54)
SECTION.bottom_margin = Cm(2.54)
SECTION.left_margin = Cm(2.5)
SECTION.right_margin = Cm(2.5)

STYLE = doc.styles['Normal']
STYLE.font.name = '宋体'
STYLE.font.size = Pt(12)
STYLE.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def _mk_run(p, text, bold=False, size=12, font='宋体', color=None, mono=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font if not mono else 'Consolas'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r

def title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _mk_run(p, text, bold=True, size=size)
    return p

def h1(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=16)
    return p

def h2(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=14)
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    _mk_run(p, text, size=12)
    return p

def body_line(parts):
    p = doc.add_paragraph()
    for t, bld, sz in parts:
        _mk_run(p, t, bold=bld, size=sz)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    _mk_run(p, text, size=9, mono=True)

def shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, h, bold=True, size=10)
        shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            _mk_run(p, str(val), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def add_img(path, caption, width=15):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, '', size=6)
        r = p.add_run()
        r.add_picture(path, width=Cm(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(cap, caption, size=9)
    return

# ===================== CONTENT =====================

doc.add_paragraph()
doc.add_paragraph()
title('《智慧水利应用》课程作业')
body('', indent=False)
title('基于大数据与机器学习的水质监测与预测系统', size=18)
body('', indent=False)
body_line([('项目进度报告（第1周）', True, 16)])
body('', indent=False)
body_line([('—— 基础框架与数据层开发', False, 14)])
doc.add_paragraph()
doc.add_paragraph()
body_line([('小组编号：第2组', False, 12)])
body_line([('小组成员：后端开发、姜玉琦、前端开发、文档统筹', False, 12)])
body_line([('提交日期：2026年5月15日', False, 12)])

doc.add_page_break()

h1('一、项目概述')
body('本项目基于Python机器学习集成方案，构建“多源数据采集→数据清洗→智能预测→可视化展示→异常告警”完整技术链路的水质监测与预测系统。系统解决传统水质监测时效性差、预测性不足、数据杂乱三大核心问题，实现水质指标的实时分析与短期预测。')
body('', indent=False)
body_line([('技术栈：', True, 12), ('Python 3.9+ / FastAPI / Pandas / Scikit-learn / XGBoost / Vue 3 / MySQL', False, 12)])
body_line([('GitHub仓库：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])

h1('二、团队分工与人员安排')
make_table(
    ['序号', '姓名', '角色', '本周主要职责', '第2周计划职责'],
    [
        ['1', '后端开发', '负责人/后端开发', '数据清洗模块开发、API服务搭建', 'ML模型开发、预测API'],
        ['2', '姜玉琦', '数据工程师', '数据采集模块（CSV/传感器/手动录入）', '数据导出功能'],
        ['3', '前端开发', '前端开发', '项目骨架搭建、单元测试编写', 'Vue 3前端框架+可视化'],
        ['4', '文档统筹', '文档/统筹', 'AI Plan文档、进度管理、PPT制作', '告警模块开发、系统集成'],
    ],
    [1.5, 2.5, 3, 5, 5]
)

doc.add_paragraph()
doc.add_page_break()

h1('三、项目进度表')
body('本系统开发周期规划为“3周开发 + 1周测试验收”，以下为总体进度规划表：')
make_table(
    ['周次', '日期', '阶段名称', '核心任务', '负责人', '交付物'],
    [
        ['第1周\n(本周)', '5/11\n~\n5/17', '基础框架\n与数据层', '项目骨架搭建\n数据采集模块\n数据清洗模块\nFastAPI基础服务', '姜玉琦\n后端开发', '项目代码框架\n采集模块API\n清洗模块API\n单元测试报告'],
        ['第2周', '5/18\n~\n5/24', '核心智能\n与展示层', 'XGBoost模型训练\n水质预测API\nVue 3前端搭建\n可视化图表', '后端开发\n前端开发', 'ML模型文件\n预测API\n前端页面'],
        ['第3周', '5/25\n~\n5/31', '功能完善\n与集成', '异常告警模块\n数据导出模块\n后台管理模块\n系统集成联调', '文档统筹\n全员', '告警功能\n导出功能\n管理界面'],
        ['第4周', '6/1\n~\n6/7', '测试与验收', '单元测试全覆盖\n集成测试\nBug修复\n验收文档', '全员', '测试报告\n验收文档'],
    ],
    [2, 2.5, 2.5, 4.5, 2.5, 3.5]
)

doc.add_paragraph()

h1('四、第1周开发完成内容')

h2('4.1 项目骨架搭建')
body('完成了完整的Python项目结构搭建，包含分层架构的7大模块目录、系统配置管理（YAML/环境变量）、Pydantic数据模型定义、FastAPI应用入口及CORS配置。')

make_table(
    ['文件路径', '功能说明', '代码行数'],
    [
        ['src/main.py', 'FastAPI应用入口（生命周期、路由注册、中间件）', '~40行'],
        ['src/config.py', '系统配置管理（数据路径、告警阈值、JWT/DB配置）', '~80行'],
        ['src/models/schemas.py', 'Pydantic数据模型（WaterQualityRecord等）', '~80行'],
        ['src/api/routes/health.py', '健康检查接口', '~15行'],
        ['src/api/routes/data_routes.py', '数据管理API（8个端点）', '~150行'],
    ],
    [5, 7, 2.5]
)

doc.add_paragraph()

h2('4.2 多源数据采集模块')
body('数据采集模块实现了三种数据源的统一接入，所有采集器继承自BaseCollector基类，统一collect()/validate()/save()接口规范。')

make_table(
    ['文件路径', '功能说明', '关键方法'],
    [
        ['src/data_collection/base.py', '采集器抽象基类', 'collect()/validate()/save()'],
        ['src/data_collection/csv_collector.py', 'CSV/Excel文件导入，支持中文列名自动映射', 'collect()/collect_batch()'],
        ['src/data_collection/sensor_collector.py', '模拟传感器数据生成，含2%异常注入', 'collect()/_generate_record()'],
        ['src/data_collection/manual_collector.py', '手动录入接口，支持单条/批量', 'collect()/collect_batch()'],
    ],
    [5, 6, 5.5]
)

doc.add_paragraph()

h2('4.3 数据清洗与管理模块')
body('实现可配置清洗流水线，基于GB 3838-2002（地表水环境质量标准）进行数据校验。清洗策略支持去重、缺失值处理（删除/插值/均值填充）、异常检测（IQR/Z-Score）、归一化（Min-Max/Z-Score）。')

make_table(
    ['文件路径', '功能说明', '关键方法'],
    [
        ['src/data_cleaning/cleaner.py', '清洗主流水线，4步处理', 'clean()-去重/缺失/异常/归一化'],
        ['src/data_cleaning/validators.py', 'GB 3838-2002标准校验器', 'validate_dataframe()/validate_record()'],
        ['src/data_cleaning/transformers.py', '时间格式统一、列名标准化、单位转换', 'standardize_datetime()/convert_units()'],
    ],
    [5, 6, 6]
)

doc.add_paragraph()

h2('4.4 FastAPI基础服务')
body('基于FastAPI框架搭建RESTful API服务，共实现8个端点。')

make_table(
    ['HTTP方法', 'API路径', '功能说明', '调用示例'],
    [
        ['GET', '/health', '系统健康检查', '返回系统状态'],
        ['POST', '/api/data/upload', '上传CSV/Excel文件', '上传540条样本'],
        ['POST', '/api/data/upload/simulate', '生成模拟传感器数据', '生成12条数据'],
        ['POST', '/api/data/manual', '手动录入水质数据', '单条录入'],
        ['GET', '/api/data/raw', '分页查询原始数据', 'page=1, page_size=10'],
        ['POST', '/api/data/clean', '执行数据清洗', '去除异常值'],
        ['GET', '/api/data/cleaned', '查询清洗后数据', '分页展示'],
        ['GET', '/api/data/summary', '获取数据统计摘要', '均值/最值/缺失统计'],
    ],
    [2, 4, 5, 5.5]
)

doc.add_paragraph()

h1('五、系统运行展示')

h2('5.1 单元测试结果')
body('所有26个单元测试全部通过（100% passed），覆盖数据采集正常流程、异常边界、数据清洗各策略、数据校验及数据转换等场景。')

make_table(
    ['测试模块', '测试用例数', '结果'],
    [
        ['数据采集-CSV导入', '4', '✅ 全部通过'],
        ['数据采集-模拟传感器', '4', '✅ 全部通过'],
        ['数据采集-手动录入', '4', '✅ 全部通过'],
        ['数据清洗-清洗流水线', '8', '✅ 全部通过'],
        ['数据清洗-数据校验', '3', '✅ 全部通过'],
        ['数据清洗-数据转换', '3', '✅ 全部通过'],
        ['合计', '26', '✅ 26/26 通过'],
    ],
    [5, 3, 4]
)

doc.add_paragraph()

h2('5.2 可视化图表展示')
chart_path = 'docs/water_quality_chart.png'
add_img(chart_path, '图1 水质监测数据可视化示例（ST001站点部分数据）')

doc.add_paragraph()

h2('5.3 API响应示例')
body_line([('示例1：健康检查 GET /health', True, 12)])
code_block('{"status": "healthy", "timestamp": "2026-05-15T12:45:57", "version": "1.0.0", "app": "Water Quality Monitoring & Prediction System"}')

body_line([('示例2：上传CSV POST /api/data/upload', True, 12)])
code_block('{"records_loaded": 540, "columns_detected": ["station_id","collection_time","ph","do","nh3n","turbidity","temperature","cod","total_phosphorus"]}')

body_line([('示例3：数据清洗 POST /api/data/clean', True, 12)])
code_block('{"total_records": 12, "duplicates_removed": 0, "outliers_removed": 2, "records_after": 10, "columns_standardized": ["ph","do","nh3n","turbidity","temperature","cod","total_phosphorus"]}')

body_line([('示例4：数据统计 GET /api/data/summary', True, 12)])
code_block('{"total_records": 540, "station_ids": ["ST001","ST002","ST003"], "date_range": ["2026-04-01","2026-04-30"], "indicators": {"ph": {"min":6.8,"max":8.14,"mean":7.51},"do": {"min":3.72,"max":7.88,"mean":5.84},"nh3n": {"min":0.01,"max":0.57,"mean":0.27}}}')

doc.add_paragraph()
doc.add_page_break()

h1('六、第2周开发规划')
body('根据项目进度表的规划，第2周将进入“核心智能与展示层”阶段，具体任务如下：')

make_table(
    ['任务编号', '任务内容', '负责人', '预计产出', '预估工时'],
    [
        ['T2.1', 'XGBoost水质预测模型训练脚本开发', '后端开发', '训练脚本、模型文件(.pkl)、验证报告', '2天'],
        ['T2.2', '特征工程（相关性分析、特征筛选）', '后端开发', '特征重要性分析报告', '1天'],
        ['T2.3', 'Vue 3 + Element Plus前端项目初始化', '前端开发', '前端项目骨架、导航页面', '1.5天'],
        ['T2.4', '水质趋势可视化（Plotly折线图、ECharts仪表盘）', '前端开发', '图表组件库', '1.5天'],
        ['T2.5', '预测API接口开发与前后端联调', '后端开发+前端开发', '联调通过的预测服务', '1天'],
        ['T2.6', '实时数据展示页面开发', '前端开发', '实时数据Web页面', '1天'],
    ],
    [1.5, 6, 2.5, 3.5, 1.5]
)

doc.add_paragraph()

h2('第2周技术方案概要')
body('ML模型方面：基于Scikit-learn + XGBoost构建水质预测模型，利田30天滑动窗口的历史数据训练，预测未来1-7天的水质指标。采用特征重要性分析筛选关键特征，网格搜索优化超参数。模型评估指标包括R²、MAE、RMSE，目标R² ≥ 0.85。')
body('前端方面：基于Vue 3 + Element Plus + ECharts搭建SPA应用，实现数据看板、趋势图表、预测结果展示等核心页面。')

doc.add_paragraph()
doc.add_page_break()

h1('七、AI辅助编程记录（AI Plan）')

h2('7.1 AI辅助编程流程')
body('本项目全程采用AI辅助编程标准流程（AI-Assisted Programming Standard Process）进行开发，开发工具为Claude Code（Claude Opus 4.7 + Sonnet 4.6）。')
body_line([('阶段1 - 需求理解与方案设计：', True, 12)])
body('AI首先阅读需求报告、概要设计和详细设计文档，理解系统功能模块划分和技术架构，然后生成项目计划和实施方案。')
body_line([('阶段2 - AI编码实施：', True, 12)])
body('AI Plan → 逐模块生成代码 → 人工Review → 测试验证 → 提交。每个模块首先定义接口规范，由AI生成完整代码实现，团队Review后测试验证，最后提交至GitHub。')
body_line([('阶段3 - 测试验证：', True, 12)])
body('AI生成测试用例 → 运行测试 → 修复Bug → 回归测试。AI根据代码逻辑自动生成覆盖正常流程和异常边界的测试用例，确保代码质量。')

doc.add_paragraph()

h2('7.2 编码指令历史记录')
make_table(
    ['指令编号', 'AI指令内容', 'AI输出文件', '状态'],
    [
        ['CMD-01', '创建FastAPI项目骨架，含配置、Pydantic模型、应用入口', 'src/config.py, main.py, models/schemas.py', '✅完成'],
        ['CMD-02', '实现多源数据采集模块（CSV导入/模拟传感器/手动录入）', 'src/data_collection/*.py (4个文件)', '✅完成'],
        ['CMD-03', '实现数据清洗流水线（去重/缺失值/异常检测/归一化）', 'src/data_cleaning/*.py (3个文件)', '✅完成'],
        ['CMD-04', '实现FastAPI路由（健康检查+数据上传/查询/清洗/统计）', 'src/api/routes/*.py (2个文件)', '✅完成'],
        ['CMD-05', '编写采集和清洗模块的单元测试', 'tests/test_collection.py, test_cleaning.py', '✅完成'],
        ['CMD-06', '生成30天3站点540条水质样本数据', 'data/samples/generate_sample.py + CSV', '✅完成'],
    ],
    [2, 6, 4.5, 2]
)

doc.add_paragraph()

h2('7.3 AI代码统计')
make_table(
    ['模块', 'Python文件数', '代码行数', '占比'],
    [
        ['项目骨架（配置+模型+入口）', '3', '~200行', '15%'],
        ['数据采集模块', '4', '~320行', '24%'],
        ['数据清洗模块', '3', '~330行', '25%'],
        ['API路由', '2', '~165行', '13%'],
        ['单元测试', '2', '~250行', '19%'],
        ['工具脚本', '1', '~50行', '4%'],
        ['合计', '15', '~1315行', '100%'],
    ],
    [4.5, 3, 3, 2.5]
)

doc.add_paragraph()

h1('八、代码框架设计与类/函数定义')
body('以下为项目核心代码的框架设计和关键类/函数定义说明：')

h2('8.1 数据采集层')
make_table(
    ['类/函数名', '所在文件', '功能描述', '输入/输出'],
    [
        ['BaseCollector', 'data_collection/base.py', '采集器抽象基类', 'collect()/validate()/save()'],
        ['CsvCollector.collect()', 'data_collection/csv_collector.py', '导入CSV/Excel并映射中文列名', 'file_path → CollectResult'],
        ['SensorCollector.collect()', 'data_collection/sensor_collector.py', '模拟生成传感器时序数据', 'station_id, hours → CollectResult'],
        ['ManualCollector.collect()', 'data_collection/manual_collector.py', '录入单条手动检测数据', 'record_dict → CollectResult'],
        ['CollectResult', 'data_collection/base.py', '采集结果数据类', 'success/records/count/errors'],
    ],
    [3.5, 4.5, 5, 4]
)

doc.add_paragraph()

h2('8.2 数据清洗层')
make_table(
    ['类/函数名', '所在文件', '功能描述', '输入/输出'],
    [
        ['DataCleaner.clean()', 'data_cleaning/cleaner.py', '执行完整清洗流水线', 'DataFrame → (DataFrame, CleaningReport)'],
        ['DataCleaner._remove_duplicates()', 'data_cleaning/cleaner.py', '基于全列去重', '内部调用'],
        ['DataCleaner._handle_missing()', 'data_cleaning/cleaner.py', '缺失值处理（drop/interpolate/fill）', '内部调用'],
        ['DataCleaner._remove_outliers()', 'data_cleaning/cleaner.py', '异常检测（IQR/zscore）', '内部调用'],
        ['WaterQualityValidator.validate_dataframe()', 'data_cleaning/validators.py', '基于GB 3838-2002标准校验', 'DataFrame → ValidationReport'],
        ['DataTransformer.standardize_datetime()', 'data_cleaning/transformers.py', '统一时间格式为ISO 8601', 'DataFrame → DataFrame'],
    ],
    [4.5, 4.5, 4.5, 4]
)

doc.add_paragraph()

h1('九、本周开发问题与解决')
make_table(
    ['序号', '问题描述', '影响范围', '解决方案', '状态'],
    [
        ['1', 'Pandas 2.0+版本兼容问题：infer_datetime_format参数已弃用', '数据转换模块', '移除已弃用参数，使用自动推断', '✅ 已解决'],
        ['2', 'Windows下中文CSV文件乱码', 'CSV导入模块', '指定utf-8-sig编码读取', '✅ 已解决'],
        ['3', '模拟传感器异常值可能极端不合理', '传感器模拟', '增加合理范围钳位逻辑', '✅ 已解决'],
        ['4', 'GitHub仓库初始化与代码推送', '项目管理', '创建remote仓库并推送', '✅ 已解决'],
    ],
    [1, 4.5, 2.5, 5, 2]
)

doc.add_paragraph()

h1('十、GitHub仓库说明')
body('本项目托管在GitHub上，小组成员通过同一代码仓库进行协同开发。')
body_line([('仓库地址：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])
body_line([('分支策略：', True, 12), ('main（主分支）+ feature/*（功能分支）', False, 12)])
body_line([('README文档：', True, 12), ('含项目说明、团队分工、项目进度表、快速开始指南、函数定义说明、API使用示例', False, 12)])
body('README中详细列出了项目的代码框架设计，包括每个程序文件的作用、每个类/函数的功能定义和高层描述，方便团队成员快速理解代码结构和协作开发。')

doc.add_paragraph()

h1('十一、下周展望')
body('第1周已经完成了系统最基础的数据层建设——数据采集和清洗功能全部就绪，API服务运行正常。第2周将聚焦核心智能能力：后端开发负责XGBoost水质预测模型的开发与训练，前端开发搭建Vue 3前端实现可视化展示，届时系统将具备完整的“数据采集→清洗→预测→展示”核心链路。')

doc.add_paragraph()
doc.add_paragraph()
body_line([('报告撰写：文档统筹', False, 12)])
body_line([('审核：后端开发', False, 12)])
body_line([('日期：2026年5月15日', False, 12)])

# Save
output = 'C:/Users/qiufengm/Desktop/智慧水利应用/第2组-水质监测预测系统-第1周进度报告.docx'
doc.save(output)
print(f'Saved to: {output}')
