"""Generate Week 4 progress report Word document."""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

SECTION = doc.sections[0]
SECTION.page_width = Emu(7560310)
SECTION.page_height = Emu(10692130)
SECTION.top_margin = Cm(2.54)
SECTION.bottom_margin = Cm(2.54)
SECTION.left_margin = Cm(2.5)
SECTION.right_margin = Cm(2.5)

STYLE = doc.styles['Normal']
STYLE.font.name = '宋体'
STYLE.font.size = Pt(12)
STYLE.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _mk_run(p, text, bold=False, size=12, font='宋体', color=None, mono=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font if not mono else 'Consolas'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r


def title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _mk_run(p, text, bold=True, size=size)
    return p


def h1(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=16)
    return p


def h2(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=14)
    return p


def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    _mk_run(p, text, size=12)
    return p


def body_line(parts):
    p = doc.add_paragraph()
    for t, bld, sz in parts:
        _mk_run(p, t, bold=bld, size=sz)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    _mk_run(p, text, size=9, mono=True)


def shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, h, bold=True, size=10)
        shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            _mk_run(p, str(val), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


# ===================== CONTENT =====================

doc.add_paragraph()
doc.add_paragraph()
title('《智慧水利应用》课程作业')
body('', indent=False)
title('基于大数据与机器学习的水质监测与预测系统', size=18)
body('', indent=False)
body_line([('项目进度报告（第4周）', True, 16)])
body('', indent=False)
body_line([('—— 测试与验收', False, 14)])
doc.add_paragraph()
doc.add_paragraph()
body_line([('小组编号：第2组', False, 12)])
body_line([('小组成员：后端开发、姜玉琦、前端开发、文档统筹', False, 12)])
body_line([('提交日期：2026年6月9日', False, 12)])

doc.add_page_break()

h1('一、项目概述')
body('本项目基于Python机器学习集成方案，构建"多源数据采集→数据清洗→智能预测→可视化展示→异常告警→基础管理"完整技术链路的水质监测与预测系统。系统解决传统水质监测时效性差、预测性不足、数据杂乱三大核心问题，实现水质指标的实时分析与短期预测。')
body('', indent=False)
body_line([('技术栈：', True, 12),
           ('Python 3.13 / FastAPI / Pandas / XGBoost / Vue 3 / Element Plus / ECharts', False, 12)])
body_line([('测试框架：', True, 12),
           ('pytest 8.4 + pytest-asyncio 0.25 + httpx 0.28', False, 12)])
body_line([('GitHub仓库：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])

h1('二、团队分工与人员安排')
make_table(
    ['序号', '姓名', '角色', '本周主要职责', '完成情况'],
    [
        ['1', '后端开发', '负责人/后端开发', '单元测试、Bug修复、后端集成', 'XPredictor/auth/alert测试 + NaN Bug修复'],
        ['2', '姜玉琦', '数据工程师', '集成测试、数据管理测试', 'DataManager测试 + CSV/Excel导出测试'],
        ['3', '前端开发', '前端开发', '前端测试、页面优化', '特征工程测试 + API集成测试'],
        ['4', '文档统筹', '文档/统筹', '验收文档、演示PPT、AI Plan', 'acceptance_report + PPT + Word报告'],
    ],
    [1.5, 2.5, 3, 5, 5.5]
)

doc.add_paragraph()
doc.add_page_break()

h1('三、项目进度表')
body('本系统开发周期为"3周开发 + 1周测试验收"，以下为总体进度规划表：')
make_table(
    ['周次', '日期', '阶段名称', '核心任务', '负责人', '交付物'],
    [
        ['第1周\n(已完成)', '5/11\n~\n5/17', '基础框架\n与数据层', '项目骨架搭建\n数据采集模块\n数据清洗模块\nFastAPI基础服务', '姜玉琦\n后端开发', '项目代码框架\n采集模块API\n清洗模块API\n单元测试报告'],
        ['第2周\n(已完成)', '5/18\n~\n5/26', '核心智能\n与展示层', 'XGBoost模型训练\n水质预测API\nVue 3前端搭建\n可视化图表', '后端开发\n前端开发', 'ML模型文件\n预测API\n前端页面'],
        ['第3周\n(已完成)', '5/27\n~\n6/2', '功能完善\n与集成', '异常告警模块\n数据导出增强\n后台管理模块\n系统集成联调', '文档统筹\n全员', '告警模块\nExcel导出\n管理后台'],
        ['第4周\n(本周)', '6/3\n~\n6/9', '测试与验收', '单元测试全覆盖\n系统集成测试\nBug修复\n验收文档', '全员', '183个测试\n验收文档\n演示PPT'],
    ],
    [2, 2.5, 2.5, 4.5, 2.5, 3.5]
)

doc.add_paragraph()

h1('四、第4周开发完成内容')

h2('4.1 单元测试全覆盖')
body('本周新增6个测试文件，156个测试用例，覆盖以下核心模块：')

make_table(
    ['测试文件', '测试数', '覆盖模块', '测试重点'],
    [
        ['test_alert_engine.py', '33', '告警引擎', 'AlertRule/Record/Engine初始化/检查/持久化/历史/规则管理'],
        ['test_auth.py', '25', '认证管理', '用户管理器/站点管理器/JWT/权限控制'],
        ['test_data_manager.py', '17', '数据管理', '原始数据/清洗数据/清空/信息查询'],
        ['test_feature_engine.py', '18', '特征工程', '滞后/滚动/差分/时间/One-Hot特征/边界情况'],
        ['test_xgboost_predictor.py', '28', 'XGBoost预测', '训练/预测/保存/加载/模型信息'],
        ['test_api_integration.py', '31', 'API全路由', '健康/数据/告警/认证/站点/导出/预测'],
    ],
    [3.5, 2, 3, 8]
)

doc.add_paragraph()

body('加上第1周的 test_collection.py（12个）和 test_cleaning.py（19个），项目总计 183个测试用例。')

doc.add_paragraph()

h2('4.2 测试运行结果')
body('全部183个测试运行通过，运行时间19.06秒，无失败、无错误。')
doc.add_paragraph()
code_block('$ python -m pytest tests/ -v')
code_block('====================== 183 passed in 19.06s ======================')

doc.add_paragraph()

body('测试验证通过的关键路径：')
body_line([('健康检查：', True, 12), ('GET /health返回status/version/app ✓', False, 12)])
body_line([('数据上传：', True, 12), ('CSV上传/模拟传感器/手动录入均正常 ✓', False, 12)])
body_line([('数据清洗：', True, 12), ('清洗报告生成、清洗数据查询 ✓', False, 12)])
body_line([('告警流程：', True, 12), ('规则获取/更新/检查/历史查询 ✓', False, 12)])
body_line([('认证流程：', True, 12), ('登录成功/失败/me/未认证 ✓', False, 12)])
body_line([('站点管理：', True, 12), ('CRUD + RBAC权限控制 ✓', False, 12)])
body_line([('数据导出：', True, 12), ('CSV/Excel/完整报告 ✓', False, 12)])
body_line([('模型预测：', True, 12), ('模型信息/训练/历史 ✓', False, 12)])

doc.add_paragraph()

h2('4.3 Bug修复')
make_table(
    ['编号', 'Bug描述', '模块', '解决方案'],
    [
        ['B-01', 'NaN JSON序列化500错误', 'API', '_safe_json辅助函数将NaN转None'],
        ['B-02', '自定义规则dict不转换', '告警引擎', '统一使用AlertRule对象传入'],
        ['B-03', '告警历史读取真实文件', '告警引擎', '测试隔离temp路径'],
        ['B-04', '测试数据跨用例污染', '认证管理', 'monkeypatch temp文件'],
        ['B-05', 'pytest-asyncio兼容性', '测试框架', '@pytest_asyncio.fixture'],
        ['B-06', 'httpx ASGITransport同步', '测试框架', '全文件async/await迁移'],
        ['B-07', '测试预期不匹配API响应', '测试', '修正assertions匹配实际响应'],
    ],
    [1.5, 5, 2.5, 7]
)

doc.add_paragraph()
doc.add_page_break()

h1('五、系统集成测试')

h2('5.1 测试架构')
body('使用 httpx 的 AsyncClient + ASGITransport 实现FastAPI全栈在进程内测试，无需启动真实服务器。测试覆盖所有6组路由、43个API端点。')
doc.add_paragraph()
code_block('transport = ASGITransport(app=app)')
code_block('async with AsyncClient(transport=transport, base_url="http://test") as client:')
code_block('    resp = await client.get("/health")')

doc.add_paragraph()

h2('5.2 测试隔离策略')
body('针对模块间状态共享问题，实施以下隔离策略：')
body_line([('告警引擎：', True, 12), ('isolated_engine fixture → temp CSV路径', False, 12)])
body_line([('认证管理：', True, 12), ('isolate_persistence fixture → monkeypatch temp JSON', False, 12)])
body_line([('DataManager：', True, 12), ('reset_data_manager fixture → 清空DataFrame', False, 12)])

doc.add_paragraph()

h1('六、系统API总览')
body('经过4周开发，系统共计43个REST API端点，覆盖6大功能模块。')
make_table(
    ['模块', '路由前缀', '端点数', '功能范围'],
    [
        ['健康检查', '/health', '2', '系统状态、应用信息'],
        ['数据管理', '/api/data', '9', '上传/查询/清洗/统计/站点/信息'],
        ['预测', '/api/predict', '5', '训练/预测/模型信息/历史'],
        ['告警', '/api/alert', '5', '规则配置/检查/历史/清空'],
        ['导出', '/api/export', '4', 'CSV/Excel/报告导出'],
        ['后台管理', '/api/admin', '8+', '登录/用户/站点CRUD'],
    ],
    [2.5, 3, 2, 8]
)

doc.add_paragraph()
doc.add_page_break()

h1('七、项目总结')

h2('7.1 功能完成情况')
body('系统经过4周迭代开发，完成了全部规划功能：')
body('1. 数据采集：CSV/Excel导入、模拟传感器、手动录入、批量录入')
body('2. 数据清洗：去重、缺失值处理、异常检测、归一化、GB3838校验')
body('3. ML预测：XGBoost 7指标独立模型、递进式多步预测、平均R²=0.825')
body('4. 告警引擎：12条GB 3838-2002规则、三级严重度、检查/历史/持久化')
body('5. 后台管理：JWT认证、RBAC权限、用户管理、站点CRUD')
body('6. 数据导出：CSV/JSON/Excel、多sheet统计报告')
body('7. 前端展示：首页看板、数据管理、预测图表、告警管理、登录、后台管理')

doc.add_paragraph()

h2('7.2 质量保障')
make_table(
    ['维度', '数据'],
    [
        ['测试用例总数', '183个'],
        ['测试通过率', '100%'],
        ['API端点', '43个'],
        ['前端页面', '6个'],
        ['ML模型', '7个（指标独立训练）'],
        ['代码总量', '~5960行'],
        ['修复Bug', '7个'],
    ],
    [4, 6]
)

doc.add_paragraph()

h2('7.3 关键技术决策与挑战')
body_line([('决策1 — XGBoost vs LSTM：', True, 12)])
body('选择XGBoost回归模型：小样本数据表现更优，训练速度快，可解释性强，适合课程项目数据规模。')
body_line([('决策2 — 文件持久化 vs 数据库：', True, 12)])
body('选择CSV/JSON文件存储：减少部署依赖，适合课程项目规模，数据可直接用Excel/文本编辑器查看。')
body_line([('决策3 — JWT无状态认证：', True, 12)])
body('选择JWT + RBAC方案：适合前后端分离架构，无服务端Session，前端token存储，简单可靠。')
body_line([('挑战1 — 特征对齐：', True, 12)])
body('训练时77维特征 vs 预测时75维特征，通过锁定feature_names + 预测时填充缺失列解决。')
body_line([('挑战2 — 测试框架兼容：', True, 12)])
body('pytest-asyncio 0.25.3 STRICT模式要求所有异步fixture和测试正确标记，通过使用@pytest_asyncio.fixture和@ pytest.mark.asyncio解决。')

doc.add_paragraph()
doc.add_page_break()

h1('八、AI辅助编程记录')

h2('8.1 AI辅助编程流程')
body('本项目全程采用AI辅助编程标准流程进行开发，开发工具为Claude Code（Claude Opus 4.7 + Sonnet 4.6）。')
body('阶段1 — 测试框架设计：AI分析各模块接口，设计测试策略，生成conftest.py共享夹具（15个fixture）。')
body('阶段2 — AI测试编码：AI Plan → 逐模块生成测试 → 运行测试 → 修复Bug → 回归测试。')
body('阶段3 — 文档生成：AI收集测试结果，生成验收报告、演示PPT大纲、Word报告。')

doc.add_paragraph()

h2('8.2 编码指令历史记录')
make_table(
    ['指令编号', 'AI指令内容', 'AI输出文件', '状态'],
    [
        ['CMD-01', '创建共享夹具（15个fixture）', 'tests/conftest.py', '✅完成'],
        ['CMD-02', '创建告警引擎测试（33个）', 'tests/test_alert_engine.py', '✅完成'],
        ['CMD-03', '创建认证管理测试（25个）', 'tests/test_auth.py', '✅完成'],
        ['CMD-04', '创建DataManager测试（17个）', 'tests/test_data_manager.py', '✅完成'],
        ['CMD-05', '创建特征工程测试（18个）', 'tests/test_feature_engine.py', '✅完成'],
        ['CMD-06', '创建XGBoost预测器测试（28个）', 'tests/test_xgboost_predictor.py', '✅完成'],
        ['CMD-07', '创建API集成测试（31个）', 'tests/test_api_integration.py', '✅完成'],
        ['CMD-08', 'Bug修复 + 测试修复（7个Bug）', 'src/api/routes/data_routes.py + tests/*', '✅完成'],
        ['CMD-09', '生成文档（AI Plan/验收/PPT/README）', 'ai_plan + docs/* + README.md', '✅完成'],
    ],
    [2, 5.5, 5, 2]
)

doc.add_paragraph()

h2('8.3 AI代码统计')
make_table(
    ['模块', '文件数', '代码行数', '占比'],
    [
        ['测试代码（6个新文件 + conftest.py）', '7', '~1860行', '~31%'],
        ['文档（AI Plan/验收/PPT/Word脚本）', '5', '~820行', '~14%'],
        ['Bug修复（data_routes.py修改）', '1', '~20行', '<1%'],
        ['Week 1-3 源码（Python + Vue）', '35+', '~4100行', '~69%'],
        ['系统总计', '~48', '~5960行', '100%'],
    ],
    [5, 2.5, 3, 2.5]
)

doc.add_paragraph()

h1('九、交付清单')
make_table(
    ['交付物', '文件/位置', '说明'],
    [
        ['源代码', 'src/ + web/', 'Python后端 ~2300行 + Vue前端 ~1800行'],
        ['测试代码', 'tests/ (8个文件)', '183个测试用例，全部通过'],
        ['AI Plan文档', 'ai_plan/week4_plan.md', '第4周AI辅助编程记录'],
        ['进度报告', 'docs/generate_week4_report.py', '第4周Word报告生成脚本'],
        ['验收文档', 'docs/acceptance_report.md', '功能验收清单 + 测试统计 + Bug清单'],
        ['演示PPT大纲', 'docs/week4_demo_ppt.md', '12页答辩PPT内容规划'],
        ['项目说明', 'README.md', '完整项目文档 + 进度表 + API文档'],
        ['示例数据', 'data/samples/', '540条水质样本数据'],
    ],
    [2.5, 5, 8]
)

doc.add_paragraph()

h1('十、GitHub仓库说明')
body('本项目托管在GitHub上，小组成员通过同一代码仓库进行协同开发。')
body_line([('仓库地址：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])
body_line([('分支策略：', True, 12), ('main（主分支）+ feature/*（功能分支）', False, 12)])
body_line([('本周提交：', True, 12), ('6个测试文件 + conftest.py + 文档 + Bug修复', False, 12)])
body('系统目前已累计43个REST API端点、6个前端页面、183个测试用例，覆盖全部6大功能模块。README文档同步更新第4周进度。')

doc.add_paragraph()

h1('十一、结语')
body('经过4周的迭代开发，基于大数据与机器学习的水质监测与预测系统完成了从需求分析、架构设计、编码实现到测试验收的全流程开发。系统实现了数据采集、清洗、预测、告警、管理、导出的完整闭环，前端可视化展示提供了直观的用户体验。')
body('在开发过程中，团队采用AI辅助编程（Claude Code）大幅提升了开发效率——测试代码的生成、Bug的定位与修复、文档的撰写均借助AI完成。同时，团队成员在各自负责的领域进行了人工审核和优化，确保了代码质量和系统可靠性。')
body('最终交付物包括：完整源代码（约5960行）、183个测试用例（全部通过）、43个API端点、6个前端页面、4周AI辅助编程记录、4份进度报告、验收文档和演示PPT。系统达到了课程设计的预期目标，具备了基本的水质监测与预测能力。')

doc.add_paragraph()
doc.add_paragraph()
body_line([('报告撰写：文档统筹', False, 12)])
body_line([('审核：后端开发', False, 12)])
body_line([('日期：2026年6月9日', False, 12)])

# Save
output = 'C:/Users/qiufengm/Desktop/智慧水利应用/第2组-水质监测预测系统-第4周进度报告.docx'
doc.save(output)
print(f'Saved to: {output}')
