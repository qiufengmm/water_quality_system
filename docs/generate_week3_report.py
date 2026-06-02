"""Generate Week 3 progress report Word document."""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

SECTION = doc.sections[0]
SECTION.page_width = Emu(7560310)
SECTION.page_height = Emu(10692130)
SECTION.top_margin = Cm(2.54)
SECTION.bottom_margin = Cm(2.54)
SECTION.left_margin = Cm(2.5)
SECTION.right_margin = Cm(2.5)

STYLE = doc.styles['Normal']
STYLE.font.name = '宋体'
STYLE.font.size = Pt(12)
STYLE.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _mk_run(p, text, bold=False, size=12, font='宋体', color=None, mono=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font if not mono else 'Consolas'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r


def title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _mk_run(p, text, bold=True, size=size)
    return p


def h1(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=16)
    return p


def h2(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=14)
    return p


def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    _mk_run(p, text, size=12)
    return p


def body_line(parts):
    p = doc.add_paragraph()
    for t, bld, sz in parts:
        _mk_run(p, t, bold=bld, size=sz)
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    _mk_run(p, text, size=9, mono=True)


def shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, h, bold=True, size=10)
        shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            _mk_run(p, str(val), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


# ===================== CONTENT =====================

doc.add_paragraph()
doc.add_paragraph()
title('《智慧水利应用》课程作业')
body('', indent=False)
title('基于大数据与机器学习的水质监测与预测系统', size=18)
body('', indent=False)
body_line([('项目进度报告（第3周）', True, 16)])
body('', indent=False)
body_line([('—— 功能完善与系统集成', False, 14)])
doc.add_paragraph()
doc.add_paragraph()
body_line([('小组编号：第2组', False, 12)])
body_line([('小组成员：后端开发、姜玉琦、前端开发、文档统筹', False, 12)])
body_line([('提交日期：2026年6月2日', False, 12)])

doc.add_page_break()

h1('一、项目概述')
body('本项目基于Python机器学习集成方案，构建"多源数据采集→数据清洗→智能预测→可视化展示→异常告警→基础管理"完整技术链路的水质监测与预测系统。系统解决传统水质监测时效性差、预测性不足、数据杂乱三大核心问题，实现水质指标的实时分析与短期预测。')
body('', indent=False)
body_line([('技术栈：', True, 12),
           ('Python 3.9+ / FastAPI / Pandas / Scikit-learn / XGBoost / Vue 3 / Element Plus / ECharts', False, 12)])
body_line([('GitHub仓库：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])

h1('二、团队分工与人员安排')
make_table(
    ['序号', '姓名', '角色', '本周主要职责', '第4周计划职责'],
    [
        ['1', '后端开发', '负责人/后端开发', '告警引擎、JWT认证、站点管理API、系统集成', '单元测试、Bug修复、验收文档'],
        ['2', '姜玉琦', '数据工程师', '数据导出增强（Excel）、集成测试', '集成测试、Bug修复'],
        ['3', '前端开发', '前端开发', '告警管理页面、登录页面、后台管理页面、前端集成', '前端告警页面优化、管理后台完善'],
        ['4', '文档统筹', '文档/统筹', 'AI Plan文档、第3周进度报告', '验收文档、演示PPT'],
    ],
    [1.5, 2.5, 3, 5.5, 5]
)

doc.add_paragraph()
doc.add_page_break()

h1('三、项目进度表')
body('本系统开发周期规划为"3周开发 + 1周测试验收"，以下为总体进度规划表：')
make_table(
    ['周次', '日期', '阶段名称', '核心任务', '负责人', '交付物'],
    [
        ['第1周\n(已完成)', '5/11\n~\n5/17', '基础框架\n与数据层', '项目骨架搭建\n数据采集模块\n数据清洗模块\nFastAPI基础服务', '姜玉琦\n后端开发', '项目代码框架\n采集模块API\n清洗模块API\n单元测试报告'],
        ['第2周\n(已完成)', '5/18\n~\n5/26', '核心智能\n与展示层', 'XGBoost模型训练\n水质预测API\nVue 3前端搭建\n可视化图表', '后端开发\n前端开发', 'ML模型文件\n预测API\n前端页面'],
        ['第3周\n(本周)', '5/27\n~\n6/2', '功能完善\n与集成', '异常告警模块\n数据导出增强\n后台管理模块\n系统集成联调', '文档统筹\n全员', '告警模块\nExcel导出\n管理后台'],
        ['第4周', '6/3\n~\n6/9', '测试与验收', '单元测试全覆盖\n集成测试\nBug修复\n验收文档', '全员', '测试报告\n验收文档\n演示PPT'],
    ],
    [2, 2.5, 2.5, 4.5, 2.5, 3.5]
)

doc.add_paragraph()

h1('四、第3周开发完成内容')

h2('4.1 异常告警模块')
body('基于GB 3838-2002《地表水环境质量标准》III类标准，实现规则驱动的告警引擎。支持12条默认规则（覆盖pH/DO/NH3N/COD/浊度/总磷6个指标），每个指标可配置运算符、阈值、严重级别（info/warning/critical）和启用状态。告警历史自动持久化到CSV文件。')
make_table(
    ['文件路径', '功能说明', '关键方法'],
    [
        ['src/alerting/alert_engine.py', '告警引擎（规则定义/数据检查/历史管理）', 'check_dataframe()/check_and_save()/get_history()'],
        ['src/api/routes/alert_routes.py', '告警API（5个端点）', 'get_rules/update_rules/check/history/clear'],
    ],
    [6, 6, 5]
)

doc.add_paragraph()

h2('4.2 告警API测试结果')
body('告警引擎功能验证：使用含超标数据的测试DataFrame进行检测，12条规则生效，正确触发10条告警（包括ST002站点DO=1.5低于2.0临界值、NH3N=3.0超标、COD=50超标等），覆盖critical/warning/info三个级别。')

doc.add_paragraph()

h2('4.3 数据导出增强')
body('在原有CSV/JSON导出基础上，新增Excel格式导出功能。使用openpyxl引擎实现.xlsx文件生成。')
make_table(
    ['API路径', '功能说明', '输出格式'],
    [
        ['GET /api/export/raw/excel', '原始数据导出', '.xlsx（单sheet）'],
        ['GET /api/export/cleaned/excel', '清洗数据导出', '.xlsx（单sheet）'],
        ['GET /api/export/report', '完整统计报告', '.xlsx（4个sheet）'],
    ],
    [5, 5, 4]
)
body('完整统计报告包含4个sheet：原始数据、清洗数据、统计摘要（含各指标均值/最值/标准差/缺失数）、数据信息（总记录数/站点列表/时间范围）。')

doc.add_paragraph()

h2('4.4 后台管理模块')
body('实现基于JWT Token的认证系统和基于角色的权限控制（RBAC），包含用户管理和站点管理功能。')
make_table(
    ['功能', 'API端点', '说明'],
    [
        ['用户登录', 'POST /api/admin/login', '用户名+密码认证，返回JWT token（2小时有效期）'],
        ['用户注册', 'POST /api/admin/register', 'admin角色专属，可创建editor/viewer用户'],
        ['用户列表', 'GET /api/admin/users', 'admin角色专属，查看所有用户信息'],
        ['当前用户', 'GET /api/admin/me', '获取当前认证用户信息'],
        ['站点列表', 'GET /api/admin/stations', '需登录，查看所有监测站点'],
        ['新增站点', 'POST /api/admin/stations', 'admin专属，添加新监测点'],
        ['更新站点', 'PUT /api/admin/stations/{id}', 'admin专属，修改站点信息'],
        ['删除站点', 'DELETE /api/admin/stations/{id}', 'admin专属，删除监测点'],
    ],
    [3, 5, 8]
)
body('用户数据通过JSON文件持久化（data/users.json），预置管理员账号admin/admin123。站点数据同样持久化为JSON文件，预置ST001上游站、ST002中游站、ST003下游站。')

doc.add_paragraph()

h2('4.5 Vue 3前端新页面')
body('本周新增3个前端页面，扩展侧边栏导航为5个菜单项。')
make_table(
    ['页面', '文件', '功能说明', '关键组件'],
    [
        ['告警管理', 'AlertManagement.vue', '规则配置（可编辑表格）、告警统计卡片、执行检查、历史记录分页', 'el-table/el-select/el-switch/el-input-number'],
        ['登录', 'Login.vue', '登录表单、演示账号提示、token存储到localStorage', 'el-form/el-input/el-card'],
        ['后台管理', 'AdminDashboard.vue', '站点CRUD（对话框表单）、用户管理、角色标识', 'el-tabs/el-dialog/el-descriptions'],
    ],
    [2.5, 4.5, 5.5, 4.5]
)
body('前端新增Axios请求拦截器自动附加Authorization头，401响应自动跳转登录页。路由守卫保护后台管理页面（/admin需登录）。侧边栏显示登录用户信息和支持退出登录。')

doc.add_paragraph()

h2('4.6 系统API总览')
body('经过3周开发，系统共计41个REST API端点，覆盖7大功能模块。')
make_table(
    ['模块', '路由前缀', '端点数', '功能范围'],
    [
        ['系统', '-', '1', '健康检查'],
        ['数据管理', '/api/data', '11', '上传/查询/清洗/统计/删除'],
        ['数据导出', '/api/export', '8', 'CSV/JSON/Excel导出、完整报告'],
        ['预测', '/api/predict', '5', '训练/预测/模型信息/历史'],
        ['告警', '/api/alert', '5', '规则配置/检查/历史'],
        ['后台管理', '/api/admin', '8', '登录/用户管理/站点CRUD'],
    ],
    [2.5, 3, 2, 8]
)

doc.add_paragraph()
doc.add_page_break()

h1('五、系统运行展示')

h2('5.1 后端验证结果')
body('全模块后端导入测试通过，41个路由注册成功。核心功能验证：')
body_line([('认证测试：', True, 12), ('POST /api/admin/login 返回JWT token ✓', False, 12)])
body_line([('站点管理：', True, 12), ('GET /api/admin/stations 返回3个站点 ✓', False, 12)])
body_line([('告警规则：', True, 12), ('GET /api/alert/rules 返回12条规则 ✓', False, 12)])
body_line([('告警检查：', True, 12), ('POST /api/alert/check 正确触发10条告警 ✓', False, 12)])
body_line([('数据导出：', True, 12), ('GET /api/export/raw/excel 下载.xlsx ✓', False, 12)])
body_line([('当前用户：', True, 12), ('GET /api/admin/me 返回管理员信息 ✓', False, 12)])

doc.add_paragraph()

h2('5.2 前端构建验证')
code_block('vite build')
code_block('✓ built in 935ms')
code_block('dist/index.html                     0.45 kB')
code_block('dist/assets/index-CpNkdR1T.css    357.52 kB')
code_block('dist/assets/index-BDMV-NL6.js   2,350.25 kB')

doc.add_paragraph()

h2('5.3 前端页面展示')
body('前端Vue 3项目构建通过，新增3个页面，侧边栏扩展为5个菜单项：')
body('1. 首页看板 - 统计卡片、站点最新数据、快速操作')
body('2. 数据管理 - CSV上传、数据清洗、表格分页浏览')
body('3. 水质预测 - 站点选择、天数滑杆、ECharts折线图')
body('4. 告警管理 - 规则编辑表格、统计卡片、告警检查、历史分页')
body('5. 后台管理 - 站点CRUD、用户管理、角色标识')
body('顶部导航栏新增登录状态显示，支持用户登录/退出操作。')

doc.add_paragraph()
doc.add_page_break()

h1('六、第4周开发规划')
body('根据项目进度表，第4周将进入"测试与验收"阶段，具体任务如下：')
make_table(
    ['任务编号', '任务内容', '负责人', '预计产出', '预估工时'],
    [
        ['T4.1', '单元测试覆盖（ML/告警/管理模块）', '前端开发', '新增测试用例', '1.5天'],
        ['T4.2', '系统集成测试与Bug修复', '全员', '集成测试报告', '2天'],
        ['T4.3', '验收文档完善', '文档统筹', '验收报告', '1天'],
        ['T4.4', '演示PPT制作', '文档统筹', '答辩PPT', '1天'],
    ],
    [1.5, 5, 2.5, 4, 1.5]
)

doc.add_paragraph()
doc.add_page_break()

h1('七、AI辅助编程记录（AI Plan）')

h2('7.1 AI辅助编程流程')
body('本项目全程采用AI辅助编程标准流程进行开发，开发工具为Claude Code（Claude Opus 4.7 + Sonnet 4.6）。')
body_line([('阶段1 - 需求理解与方案设计：', True, 12)])
body('AI首先阅读第2周代码结构和现有API，分析系统现有路由注册模式和前端组件设计风格，设计告警引擎架构和JWT认证方案。')
body_line([('阶段2 - AI编码实施：', True, 12)])
body('AI Plan → 逐模块生成代码 → 人工Review → 测试验证 → 提交。告警模块先完成引擎核心，再实现API路由。后台管理先完成JWT认证，再实现站点CRUD。前端按"页面创建→路由注册→API集成"顺序推进。')
body_line([('阶段3 - 测试验证：', True, 12)])
body('AI生成集成测试 → 运行测试 → 修复依赖问题（bcrypt版本兼容性）→ 回归测试。')

doc.add_paragraph()

h2('7.2 编码指令历史记录')
make_table(
    ['指令编号', 'AI指令内容', 'AI输出文件', '状态'],
    [
        ['CMD-01', '创建告警引擎（AlertRule/AlertRecord/AlertEngine + 默认规则）', 'src/alerting/*.py (2个文件)', '✅完成'],
        ['CMD-02', '实现告警API路由（规则/检查/历史/清空）', 'src/api/routes/alert_routes.py', '✅完成'],
        ['CMD-03', '实现JWT认证和用户管理（User/UserManager/Token）', 'src/admin/auth.py', '✅完成'],
        ['CMD-04', '实现站点管理和后台API（站点CRUD + admin_routes）', 'src/admin/station_manager.py + admin_routes.py', '✅完成'],
        ['CMD-05', '增强数据导出（Excel格式 + 完整报告）', 'src/api/routes/export_routes.py (修改)', '✅完成'],
        ['CMD-06', '实现3个前端页面（告警管理/登录/后台管理）', 'web/src/views/*.vue (3个文件)', '✅完成'],
        ['CMD-07', '前端集成（路由/Axios拦截器/App.vue侧边栏+登录状态）', 'web/src/* (3个文件修改)', '✅完成'],
    ],
    [2, 6, 5, 2]
)

doc.add_paragraph()

h2('7.3 AI代码统计')
make_table(
    ['模块', '文件数', '代码行数', '占比'],
    [
        ['告警引擎 + API', '3', '~275行', '17%'],
        ['后台管理（认证+站点+API）', '4', '~435行', '27%'],
        ['数据导出增强', '1', '~80行(修改)', '5%'],
        ['前端新页面（告警/登录/后台管理）', '3', '~680行', '43%'],
        ['前端集成（API/路由/App.vue）', '3', '~110行(修改)', '7%'],
        ['合计', '14', '~1590行', '100%'],
    ],
    [4.5, 2.5, 3, 2.5]
)

doc.add_paragraph()

h1('八、代码框架设计与类/函数定义')
body('以下为第3周新增的核心代码框架和关键类/函数定义说明：')

h2('8.1 告警引擎')
make_table(
    ['类/函数名', '所在文件', '功能描述', '输入/输出'],
    [
        ['AlertRule', 'alerting/alert_engine.py', '告警规则数据类：indicator/operator/threshold/severity/enabled', '数据类'],
        ['AlertRecord', 'alerting/alert_engine.py', '告警记录数据类：station_id/indicator/value/rule/timestamp/status', '数据类'],
        ['AlertEngine.check_dataframe()', 'alerting/alert_engine.py', '检查DataFrame中所有记录触发的告警', 'DataFrame → list[AlertRecord]'],
        ['AlertEngine.check_and_save()', 'alerting/alert_engine.py', '检查并持久化新告警记录到CSV', 'DataFrame → list[AlertRecord]'],
        ['AlertEngine.get_history()', 'alerting/alert_engine.py', '分页查询告警历史（支持按级别过滤）', 'page, page_size → dict'],
    ],
    [4.5, 5, 5, 4.5]
)

doc.add_paragraph()

h2('8.2 后台管理')
make_table(
    ['类/函数名', '所在文件', '功能描述', '输入/输出'],
    [
        ['User', 'admin/auth.py', '用户数据类：username/password_hash/role/display_name', '数据类'],
        ['UserManager.authenticate()', 'admin/auth.py', '验证用户名密码', 'username, password → Optional[User]'],
        ['UserManager.create_user()', 'admin/auth.py', '创建新用户（带密码哈希）', 'username, password, role → bool'],
        ['Station', 'admin/auth.py', '站点数据类：station_id/name/location/description/contact', '数据类'],
        ['StationManager.add_station()', 'admin/auth.py', '新增监测站点', 'Station → bool'],
        ['StationManager.update_station()', 'admin/auth.py', '更新站点信息', 'station_id, data → bool'],
        ['create_access_token()', 'admin/auth.py', '创建JWT访问令牌', 'data → str(token)'],
        ['verify_token()', 'admin/auth.py', '验证JWT令牌', 'token → Optional[dict]'],
        ['get_current_user()', 'admin/auth.py', 'FastAPI依赖注入（从token获取用户）', 'credentials → User'],
        ['require_role()', 'admin/auth.py', '角色权限检查依赖工厂', 'role → dependency'],
    ],
    [4.5, 5, 5, 4.5]
)

doc.add_paragraph()

h2('8.3 新增前端组件')
make_table(
    ['组件名', '文件路径', '功能描述', '关键依赖'],
    [
        ['AlertManagement.vue', 'web/src/views/AlertManagement.vue', '告警规则编辑、统计卡片、执行检查、历史分页', 'el-table/el-switch/el-input-number'],
        ['Login.vue', 'web/src/views/Login.vue', '登录表单、token存储、路由跳转', 'el-form/el-input/axios'],
        ['AdminDashboard.vue', 'web/src/views/AdminDashboard.vue', '站点CRUD对话框、用户管理表格、角色标签', 'el-tabs/el-dialog/el-descriptions'],
    ],
    [3, 5, 5.5, 5]
)

doc.add_paragraph()

h1('九、本周开发问题与解决')
make_table(
    ['序号', '问题描述', '影响范围', '解决方案', '状态'],
    [
        ['1', 'passlib + bcrypt版本不兼容（__about__属性缺失）', '用户认证模块', '降级bcrypt到4.0.1版本', '已解决'],
        ['2', 'FastAPI Query的regex参数弃用警告', '告警API', '替换regex为pattern参数', '已解决'],
    ],
    [1, 5, 3, 5, 2]
)

doc.add_paragraph()

h1('十、GitHub仓库说明')
body('本项目托管在GitHub上，小组成员通过同一代码仓库进行协同开发。')
body_line([('仓库地址：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])
body_line([('分支策略：', True, 12), ('main（主分支）+ feature/*（功能分支）', False, 12)])
body_line([('本周提交：', True, 12), ('告警模块 + 后台管理 + Excel导出增强 + 3个前端页面', False, 12)])
body('系统目前已累计41个REST API端点、18个前端页面组件，覆盖全部7个功能模块。README文档同步更新第3周进度。')

doc.add_paragraph()

h1('十一、下周展望')
body('第3周完成了系统的功能完善——异常告警模块（GB 3838-2002标准12条规则）、数据导出增强（Excel多sheet报告）、后台管理（JWT认证+RBAC+站点CRUD）。系统已具备完整的"数据采集→清洗→预测→告警→管理→导出"全链路能力。')
body('第4周将进入测试验收阶段：前端开发负责新增模块的单元测试，全员进行集成测试和Bug修复，文档统筹完善验收文档和答辩PPT。届时系统将完成课程要求的全部七个功能模块，具备完整的验收交付条件。')

doc.add_paragraph()
doc.add_paragraph()
body_line([('报告撰写：文档统筹', False, 12)])
body_line([('审核：后端开发', False, 12)])
body_line([('日期：2026年6月2日', False, 12)])

# Save
output = 'C:/Users/qiufengm/Desktop/智慧水利应用/第2组-水质监测预测系统-第3周进度报告.docx'
doc.save(output)
print(f'Saved to: {output}')
