"""Generate Week 2 progress report Word document."""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

SECTION = doc.sections[0]
SECTION.page_width = Emu(7560310)
SECTION.page_height = Emu(10692130)
SECTION.top_margin = Cm(2.54)
SECTION.bottom_margin = Cm(2.54)
SECTION.left_margin = Cm(2.5)
SECTION.right_margin = Cm(2.5)

STYLE = doc.styles['Normal']
STYLE.font.name = '宋体'
STYLE.font.size = Pt(12)
STYLE.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def _mk_run(p, text, bold=False, size=12, font='宋体', color=None, mono=False):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = font if not mono else 'Consolas'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r

def title(text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _mk_run(p, text, bold=True, size=size)
    return p

def h1(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=16)
    return p

def h2(text):
    p = doc.add_paragraph()
    _mk_run(p, text, bold=True, size=14)
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    _mk_run(p, text, size=12)
    return p

def body_line(parts):
    p = doc.add_paragraph()
    for t, bld, sz in parts:
        _mk_run(p, t, bold=bld, size=sz)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    _mk_run(p, text, size=9, mono=True)

def shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, h, bold=True, size=10)
        shade(c, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            _mk_run(p, str(val), size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def add_img(path, caption, width=15):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(p, '', size=6)
        r = p.add_run()
        r.add_picture(path, width=Cm(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _mk_run(cap, caption, size=9)
    return

# ===================== CONTENT =====================

doc.add_paragraph()
doc.add_paragraph()
title('《智慧水利应用》课程作业')
body('', indent=False)
title('基于大数据与机器学习的水质监测与预测系统', size=18)
body('', indent=False)
body_line([('项目进度报告（第2周）', True, 16)])
body('', indent=False)
body_line([('—— 核心智能与展示层开发', False, 14)])
doc.add_paragraph()
doc.add_paragraph()
body_line([('小组编号：第2组', False, 12)])
body_line([('小组成员：谢坤、姜玉琦、苏航、赵宏斌', False, 12)])
body_line([('提交日期：2026年5月26日', False, 12)])

doc.add_page_break()

h1('一、项目概述')
body('本项目基于Python机器学习集成方案，构建"多源数据采集→数据清洗→智能预测→可视化展示→异常告警"完整技术链路的水质监测与预测系统。系统解决传统水质监测时效性差、预测性不足、数据杂乱三大核心问题，实现水质指标的实时分析与短期预测。')
body('', indent=False)
body_line([('技术栈：', True, 12), ('Python 3.9+ / FastAPI / Pandas / Scikit-learn / XGBoost / Vue 3 / Element Plus / ECharts', False, 12)])
body_line([('GitHub仓库：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])

h1('二、团队分工与人员安排')
make_table(
    ['序号', '姓名', '角色', '本周主要职责', '第3周计划职责'],
    [
        ['1', '谢坤', '负责人/后端开发', 'ML模型开发（XGBoost预测器+特征工程+训练脚本）、预测API', '后台管理模块（JWT认证、RBAC权限）'],
        ['2', '姜玉琦', '数据工程师', '数据清洗优化、集成测试', '数据导出功能（Excel/PDF报表）'],
        ['3', '苏航', '前端开发', 'Vue 3 + Element Plus前端搭建、ECharts可视化图表、3个核心页面', '前端告警页面、管理后台页面'],
        ['4', '赵宏斌', '文档/统筹', 'AI Plan文档、第2周进度报告、PPT制作', '告警模块开发、系统集成'],
    ],
    [1.5, 2.5, 3, 5.5, 5]
)

doc.add_paragraph()
doc.add_page_break()

h1('三、项目进度表')
body('本系统开发周期规划为"3周开发 + 1周测试验收"，以下为总体进度规划表：')
make_table(
    ['周次', '日期', '阶段名称', '核心任务', '负责人', '交付物'],
    [
        ['第1周\n(已完成)', '5/11\n~\n5/17', '基础框架\n与数据层', '项目骨架搭建\n数据采集模块\n数据清洗模块\nFastAPI基础服务', '姜玉琦\n谢坤', '项目代码框架\n采集模块API\n清洗模块API\n单元测试报告'],
        ['第2周\n(本周)', '5/18\n~\n5/26', '核心智能\n与展示层', 'XGBoost模型训练\n水质预测API\nVue 3前端搭建\n可视化图表', '谢坤\n苏航', 'ML模型文件\n预测API\n前端页面'],
        ['第3周', '5/27\n~\n6/2', '功能完善\n与集成', '异常告警模块\n数据导出模块\n后台管理模块\n系统集成联调', '赵宏斌\n全员', '告警功能\n导出功能\n管理界面'],
        ['第4周', '6/3\n~\n6/9', '测试与验收', '单元测试全覆盖\n集成测试\nBug修复\n验收文档', '全员', '测试报告\n验收文档'],
    ],
    [2, 2.5, 2.5, 4.5, 2.5, 3.5]
)

doc.add_paragraph()

h1('四、第2周开发完成内容')

h2('4.1 ML模型模块（XGBoost水质预测器）')
body('基于XGBoost回归算法，为每个水质指标训练独立的预测模型。特征工程包括：滞后特征（t-1至t-7）、滚动窗口统计（3期均值/标准差）、一阶差分、时间特征（小时/天/月/星期）、站点One-Hot编码，共生成77维特征向量。')

make_table(
    ['文件路径', '功能说明', '关键方法'],
    [
        ['src/ml/base.py', '预测器抽象基类 + 预测结果数据类', 'train()/predict()/save_model()/load_model()'],
        ['src/ml/feature_engineer.py', '特征工程（滞后/滚动/差分/时间/One-Hot）', 'create_features()/create_prediction_features()'],
        ['src/ml/xgboost_predictor.py', 'XGBoost预测器实现', 'train()/predict()/save_model()/load_model()'],
        ['src/ml/train.py', '训练脚本（加载→清洗→特征→训练→评估→保存）', 'train_model()/main()'],
    ],
    [5, 6, 6]
)

doc.add_paragraph()

h2('4.2 模型训练结果')
body('使用540条样本数据（3站点，30天，每4小时采集），清洗后保留525条有效记录，训练504条（80%训练集/20%验证集）。平均R²=0.8245，77维特征，200棵决策树。')

make_table(
    ['水质指标', 'R²', 'MAE', 'RMSE', '评价'],
    [
        ['pH', '0.8316', '0.055', '0.068', '✅ 优秀'],
        ['溶解氧(DO)', '0.7040', '0.170', '0.211', '✅ 良好'],
        ['氨氮(NH3N)', '0.5576', '0.021', '0.027', '⚠ 一般'],
        ['浊度(Turbidity)', '0.9377', '0.125', '0.169', '✅ 优秀'],
        ['水温(Temperature)', '0.9155', '0.530', '0.735', '✅ 优秀'],
        ['化学需氧量(COD)', '0.9362', '0.254', '0.353', '✅ 优秀'],
        ['总磷(Total_P)', '0.8890', '0.004', '0.006', '✅ 优秀'],
    ],
    [3.5, 2, 2, 2, 2.5]
)
body('多数指标R² > 0.83，模型具有较好的预测能力。NH3N指标R²=0.56偏低，可能因数据中NH3N变化规律较复杂，后续可增加更多相关特征以提升精度。', indent=True)

doc.add_paragraph()

h2('4.3 预测API接口')
body('实现5个预测相关的REST API端点，完成从模型训练到预测查询的完整链路。')

make_table(
    ['HTTP方法', 'API路径', '功能说明'],
    [
        ['POST', '/api/predict/train', '基于样本数据训练XGBoost模型'],
        ['POST', '/api/predict/train/from-data', '基于已加载数据训练（自动清洗前置）'],
        ['POST', '/api/predict/batch', '指定站点+天数进行批量预测'],
        ['GET', '/api/predict/model-info', '查询当前模型信息（指标/特征数/R²）'],
        ['GET', '/api/predict/history', '查看历史训练模型列表'],
    ],
    [2, 5, 9]
)

doc.add_paragraph()

h2('4.4 Vue 3前端页面')
body('基于Vue 3 + Element Plus + ECharts搭建SPA应用，实现3个核心页面，通过Axios代理访问FastAPI后端。')

make_table(
    ['页面', '文件', '功能说明', '关键组件'],
    [
        ['首页看板', 'Dashboard.vue', '统计卡片（总记录/站点/模型状态）、站点数据卡片、快速操作入口', 'el-card/el-row/el-tag'],
        ['数据管理', 'DataManagement.vue', 'CSV拖拽上传、数据清洗执行、清洗报告、数据表格分页浏览', 'el-upload/el-table/el-dialog'],
        ['水质预测', 'Prediction.vue', '站点选择、天数滑杆、模型训练、ECharts折线图展示预测结果、详情表格', 'el-select/el-slider/echarts'],
    ],
    [2.5, 4, 6, 4]
)

doc.add_paragraph()

h2('4.5 Bug修复总结')
body('本周开发过程中发现并修复了3个关键Bug，保障了系统的稳定运行。')

make_table(
    ['序号', '问题', '根因', '修复方案', '状态'],
    [
        ['1', '滚动窗口统计TypeError', '.transform(lambda x: x.rolling(...))返回SeriesGroupBy，.mean()调用失败', '在transform内部直接完成rolling().mean()和rolling().std()', '已解决'],
        ['2', '预测特征列数不匹配', '训练时3站点生成77列，预测时1站点生成75列，XGBoost报错', '锁定feature_names + 预测时补全缺失列（赋值为0）', '已解决'],
        ['3', '训练数据含NaN导致失败', '原始数据直接传入训练，未清洗导致标签含NaN', '训练前增加DataCleaner前置清洗步骤', '已解决'],
    ],
    [1, 3.5, 4.5, 5, 2]
)

doc.add_paragraph()
doc.add_page_break()

h1('五、系统运行展示')

h2('5.1 API测试结果')
body('全流程集成测试全部通过：上传540条CSV数据 → 清洗（540→525条） → 训练XGBoost（R²=0.8245） → 预测（7天7指标） → 模型信息查询。')

body_line([('示例1：训练模型 POST /api/predict/train/from-data', True, 12)])
code_block('{"message": "Model trained on loaded data", "records_used": 504, "avg_r2": 0.8245}')

body_line([('示例2：批量预测 POST /api/predict/batch?station_id=ST001&days=7', True, 12)])
code_block('{"station_id": "ST001", "prediction_dates": ["2026-04-30",..."2026-05-06"], "confidence": 0.8245}')

body_line([('示例3：模型信息 GET /api/predict/model-info', True, 12)])
code_block('{"is_trained": true, "target_indicators": ["ph","do","nh3n","turbidity","temperature","cod","total_phosphorus"], "num_features": 77}')

doc.add_paragraph()

h2('5.2 前端页面展示')
body('Vue 3前端项目通过Vite构建，代码编译通过，生成3个功能页面：')
body('1. 首页看板：显示总记录数（540）、站点数（3）、模型状态、各站点最新水质数据及等级标识（I~II类/III类/IV~V类），提供快速操作入口。')
body('2. 数据管理：支持CSV文件拖拽上传、模拟数据生成、一键数据清洗（含清洗报告弹窗）、原始/清洗数据分页表格浏览。')
body('3. 水质预测：选择站点和预测天数（1-30天），一键训练模型，ECharts折线图展示7个指标的未来趋势预测。')

doc.add_paragraph()

h2('5.3 前端构建验证')
code_block('vite build')
code_block('✓ built in 941ms')
code_block('dist/index.html                     0.45 kB')
code_block('dist/assets/index-BjOVE5_R.css    356.98 kB')
code_block('dist/assets/index-CgIto4IL.js   2,329.57 kB')

doc.add_paragraph()
doc.add_page_break()

h1('六、第3周开发规划')
body('根据项目进度表，第3周将进入"功能完善与集成"阶段，具体任务如下：')

make_table(
    ['任务编号', '任务内容', '负责人', '预计产出', '预估工时'],
    [
        ['T3.1', '异常告警模块（阈值配置、自动触发、历史记录）', '赵宏斌', '告警引擎、告警历史API', '2天'],
        ['T3.2', '数据导出功能（Excel/PDF报表）', '姜宇琦', '导出API、统计报表', '1.5天'],
        ['T3.3', '后台管理（JWT认证、RBAC权限、站点管理）', '谢坤', '登录注册、权限控制、点位管理', '2天'],
        ['T3.4', '系统集成联调（前后端全流程打通）', '全员', '集成测试报告', '1天'],
    ],
    [1.5, 5, 2.5, 4, 1.5]
)

doc.add_paragraph()

h2('第3周技术方案概要')
body('告警模块基于GB 3838-2002标准配置水质指标阈值，支持自动触发告警并记录历史。数据导出支持Excel和PDF格式的统计报表。后台管理采用JWT Token认证，基于角色的权限控制（RBAC），包含站点管理和系统配置功能。')

doc.add_paragraph()
doc.add_page_break()

h1('七、AI辅助编程记录（AI Plan）')

h2('7.1 AI辅助编程流程')
body('本项目全程采用AI辅助编程标准流程进行开发，开发工具为Claude Code（Claude Opus 4.7 + Sonnet 4.6）。')
body_line([('阶段1 - 需求理解与方案设计：', True, 12)])
body('AI首先阅读第1周代码结构和需求文档，分析现有系统功能模块和数据流，设计ML模型架构和前端页面方案。')
body_line([('阶段2 - AI编码实施：', True, 12)])
body('AI Plan → 逐模块生成代码 → 人工Review → 测试验证 → 提交。ML模块先完成基础框架（Base + FeatureEngineer + XGBoostPredictor），再实现训练脚本和预测API。前端通过Vite模板初始化后，AI生成3个核心页面代码。')
body_line([('阶段3 - 测试验证：', True, 12)])
body('AI生成集成测试 → 运行测试 → 修复Bug → 回归测试。本周发现并修复了3个关键Bug，均通过AI诊断根因并生成修复方案。')

doc.add_paragraph()

h2('7.2 编码指令历史记录')
make_table(
    ['指令编号', 'AI指令内容', 'AI输出文件', '状态'],
    [
        ['CMD-01', '创建ML基础框架（Base + FeatureEngineer + XGBoostPredictor + Train）', 'src/ml/*.py (4个文件)', '✅完成'],
        ['CMD-02', '实现预测API路由（训练/预测/模型信息/历史）', 'src/api/routes/predict_routes.py', '✅完成'],
        ['CMD-03', '创建Vue 3 + Vite前端项目，安装Element Plus/ECharts/Axios', 'web/项目框架', '✅完成'],
        ['CMD-04', '实现3个前端页面（Dashboard/DataManagement/Prediction）', 'web/src/views/*.vue (3个文件)', '✅完成'],
        ['CMD-05', '修复滚动窗口Bug + 特征列对齐Bug + 训练NaN Bug', 'feature_engineer.py/xgboost_predictor.py/predict_routes.py', '✅完成'],
    ],
    [2, 6, 5, 2]
)

doc.add_paragraph()

h2('7.3 AI代码统计')
make_table(
    ['模块', '文件数', '代码行数', '占比'],
    [
        ['ML模块（Base + FeatureEngineer + XGBoostPredictor + Train）', '4', '~840行', '40%'],
        ['预测API路由', '1', '~180行', '9%'],
        ['前端Vue项目配置 + 入口', '3', '~105行', '5%'],
        ['前端页面（Dashboard + DataManagement + Prediction）', '3', '~600行', '29%'],
        ['前端HTTP封装 + 路由', '2', '~100行', '5%'],
        ['Bug修复', '3', '~15行(修改)', '1%'],
        ['文档（AI Plan + Word报告生成）', '2', '~230行', '11%'],
        ['合计', '18', '~2070行', '100%'],
    ],
    [4.5, 2.5, 3, 2.5]
)

doc.add_paragraph()

h1('八、代码框架设计与类/函数定义')
body('以下为第2周新增的核心代码框架和关键类/函数定义说明：')

h2('8.1 ML预测层')
make_table(
    ['类/函数名', '所在文件', '功能描述', '输入/输出'],
    [
        ['AbstractPredictor', 'ml/base.py', '预测器抽象基类，定义ML接口规范', 'train()/predict()/save_model()/load_model()'],
        ['PredictionResult', 'ml/base.py', '预测结果数据类', 'success/station_id/predictions/dates/confidence'],
        ['FeatureEngineer.create_features()', 'ml/feature_engineer.py', '从时序数据创建特征矩阵（77维）', 'DataFrame → DataFrame(特征矩阵)'],
        ['FeatureEngineer.create_prediction_features()', 'ml/feature_engineer.py', '创建预测特征（不含目标列）', 'DataFrame → DataFrame(推理特征)'],
        ['XGBoostPredictor.train()', 'ml/xgboost_predictor.py', '训练每个指标的XGBoost回归模型', 'DataFrame → metrics dict'],
        ['XGBoostPredictor.predict()', 'ml/xgboost_predictor.py', '递进式多步预测未来水质', 'DataFrame, days → PredictionResult'],
        ['XGBoostPredictor.save_model()', 'ml/xgboost_predictor.py', '保存模型+元数据+特征配置到磁盘', 'path → model_dir_path'],
        ['XGBoostPredictor.load_model()', 'ml/xgboost_predictor.py', '从磁盘加载已训练的模型', 'path → bool'],
        ['train_model()', 'ml/train.py', '完整训练流程（加载→清洗→特征→训练→评估→保存）', 'data_path, ... → XGBoostPredictor'],
    ],
    [4.5, 4, 5, 4.5]
)

doc.add_paragraph()

h2('8.2 预测API层')
make_table(
    ['函数名', 'API路径', '功能描述', '输入/输出'],
    [
        ['train_prediction_model()', 'POST /api/predict/train', '基于样本数据训练模型', '无参 → train结果'],
        ['train_from_loaded_data()', 'POST /api/predict/train/from-data', '基于已加载数据清洗后训练', '无参 → train结果含R²'],
        ['predict_batch()', 'POST /api/predict/batch', '指定站点+天数批量预测', 'station_id, days → 预测值+置信度'],
        ['get_model_info()', 'GET /api/predict/model-info', '查询当前模型信息', '无参 → 模型配置+指标+特征重要性'],
        ['get_prediction_history()', 'GET /api/predict/history', '查看历史保存的模型', '无参 → 模型列表'],
    ],
    [3.5, 4.5, 5, 4.5]
)

doc.add_paragraph()

h2('8.3 前端Vue组件')
make_table(
    ['组件名', '文件路径', '功能描述', '关键依赖'],
    [
        ['App.vue', 'web/src/App.vue', '根布局（侧边栏导航+顶部标题+内容区）', 'Element Plus Container/Menu'],
        ['Dashboard.vue', 'web/src/views/Dashboard.vue', '首页看板（统计卡片+站点数据+快速操作）', 'el-card/el-row/el-tag'],
        ['DataManagement.vue', 'web/src/views/DataManagement.vue', '数据管理（上传+清洗+表格浏览）', 'el-upload/el-table/el-dialog'],
        ['Prediction.vue', 'web/src/views/Prediction.vue', '水质预测（控制面板+ECharts图表+详情表）', 'echarts/el-slider/el-select'],
        ['api/index.js', 'web/src/api/index.js', 'Axios HTTP封装（15个API接口）', 'axios'],
        ['router/index.js', 'web/src/router/index.js', 'Vue Router路由配置', 'vue-router'],
    ],
    [3, 4.5, 4, 4.5]
)

doc.add_paragraph()

h1('九、本周开发问题与解决')
make_table(
    ['序号', '问题描述', '影响范围', '解决方案', '状态'],
    [
        ['1', 'FeatureEngineer滚动窗口统计：transform返回SeriesGroupBy调用mean()失败', 'ML特征工程', '在transform内部直接完成rolling聚合操作', '已解决'],
        ['2', '预测特征列数不匹配：训练77列vs预测75列（One-Hot编码差异）', 'XGBoost预测', '锁定训练特征列名，预测时自动补全缺失列', '已解决'],
        ['3', '训练数据含NaN：原始数据未清洗直接传入导致XGBoost标签验证失败', '预测API', '增加DataCleaner前置清洗步骤', '已解决'],
        ['4', 'Windows GBK编码：控制台打印R²符号报UnicodeEncodeError', '训练脚本', '用R^2替代R²符号', '已解决'],
        ['5', 'Vue前端构建：WaterLevel图标不存在、do是JS保留关键字', '前端页面', '替换为Monitor图标、参数名改为doVal', '已解决'],
    ],
    [1, 4.5, 2.5, 5, 2]
)

doc.add_paragraph()

h1('十、GitHub仓库说明')
body('本项目托管在GitHub上，小组成员通过同一代码仓库进行协同开发。')
body_line([('仓库地址：', True, 12), ('https://github.com/qiufengmm/water_quality_system', False, 12)])
body_line([('分支策略：', True, 12), ('main（主分支）+ feature/*（功能分支）', False, 12)])
body_line([('本周提交：', True, 12), ('ML模块 + 预测API + Vue 3前端 + 可视化页面 + Bug修复', False, 12)])
body('README文档中更新了第2周进度表、ML模块函数定义、前端模块说明、以及快速启动步骤（含前端启动命令）。')

doc.add_paragraph()

h1('十一、下周展望')
body('第2周完成了系统的核心智能能力——XGBoost水质预测模型（R²=0.8245）和Vue 3前端可视化页面。系统现具备"数据采集→清洗→预测→展示"的完整核心链路。')
body('第3周将聚焦功能完善：赵宏斌负责异常告警模块（基于GB 3838-2002标准）、姜玉琦负责数据导出功能、谢坤负责后台管理JWT认证和RBAC权限，全员集成联调，为第4周测试验收做准备。')
body('届时，系统将覆盖课程要求的全部七个功能模块，形成完整的水质监测与预测解决方案。')

doc.add_paragraph()
doc.add_paragraph()
body_line([('报告撰写：赵宏斌', False, 12)])
body_line([('审核：谢坤', False, 12)])
body_line([('日期：2026年5月26日', False, 12)])

# Save
output = 'C:/Users/qiufengm/Desktop/智慧水利应用/第2组-水质监测预测系统-第2周进度报告.docx'
doc.save(output)
print(f'Saved to: {output}')
